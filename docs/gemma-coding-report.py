#!/usr/bin/env python3
"""Generate the Gemma coding report as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(18)
style_h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

style_h2 = doc.styles['Heading 2']
style_h2.font.size = Pt(14)
style_h2.font.color.rgb = RGBColor(0x2d, 0x2d, 0x5e)

style_h3 = doc.styles['Heading 3']
style_h3.font.size = Pt(12)

# --- Title Page ---
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Getting Small AI Models to Code')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Technical Report on Engineering Gemma 4 E2B\nfrom 0% to 65% on Multi-Function Coding Benchmarks')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x88)

doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Russell Hanson\nApril 2026\n\nModel: Google Gemma 4 E2B (27B parameters)\nPlatform: LM Studio on Mac Metal GPU\nFramework: COD (Open Source Claude Code Implementation)')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# --- Table of Contents ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1.', 'Executive Summary', '3'),
    ('2.', 'The Challenge: Why Small Models Struggle with Code', '4'),
    ('3.', 'Baseline Assessment: Where Gemma Starts', '5'),
    ('4.', 'Infrastructure Problems: Death by a Thousand Cuts', '7'),
    ('5.', 'The Streaming Problem: Why Thinking Models Need Special Handling', '9'),
    ('6.', 'System Prompt Engineering: Algorithmic Hints That Work', '11'),
    ('7.', 'Self-Correction Loops: Teaching Models to Fix Their Own Bugs', '13'),
    ('8.', 'What Didn\'t Work: Failed Approaches', '15'),
    ('9.', 'Final Architecture and Results', '17'),
    ('10.', 'Lessons Learned: A Primer for Small Model Coding', '19'),
]
for num, title_text, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title_text}')
    run.font.size = Pt(11)

doc.add_page_break()

# --- Chapter 1: Executive Summary ---
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'This report documents the engineering effort required to make Google\'s Gemma 4 E2B '
    '(a 27-billion parameter language model running locally on a Mac with Metal GPU acceleration) '
    'perform competently on a multi-function coding benchmark. The benchmark requires implementing '
    'five distinct algorithms in a single Python file: an LRU Cache with TTL, text justification, '
    'interval painting, an expression evaluator with operator precedence, and a Roman numeral calculator.'
)

doc.add_paragraph(
    'The journey from 0% to 65% (32/49 test cases passing) required solving problems across '
    'four distinct layers: infrastructure bugs in the agent framework, streaming API incompatibilities '
    'with thinking models, system prompt engineering with algorithmic hints, and multi-turn '
    'self-correction loops. No single change was sufficient; the final score required all four layers working together.'
)

doc.add_heading('Score Progression', level=3)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Phase', 'Score', 'Key Change']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ('Initial Attempt', '0/49 (0%)', 'CLI flag parsing bug prevented Gemma from being selected'),
    ('First Successful Run', '21/49 (43%)', 'Fixed CLI, added reasoning_effort, increased max_tokens'),
    ('After Iterative Strategy', '0/33 (0%)', 'Iterative approach broke file assembly (regression)'),
    ('After System Prompt Hints', '19/41 (46%)', 'Algorithmic hints improved roman_calc and evaluate_v2'),
    ('After Bug Fixes + Self-Correction', '32/49 (65%)', 'Fixed LRU double-delete, justify logic, non-streaming adapter'),
]
for row_idx, (phase, score, change) in enumerate(data, 1):
    table.rows[row_idx].cells[0].text = phase
    table.rows[row_idx].cells[1].text = score
    table.rows[row_idx].cells[2].text = change

doc.add_paragraph('')
doc.add_paragraph(
    'The central lesson is that getting small models to code is not primarily a model capability problem\u2014'
    'it is an infrastructure and prompting problem. The model has the knowledge to implement these algorithms; '
    'the challenge is creating an environment where it can express that knowledge reliably.'
)

doc.add_page_break()

# --- Chapter 2: The Challenge ---
doc.add_heading('2. The Challenge: Why Small Models Struggle with Code', level=1)

doc.add_paragraph(
    'Large language models like Claude Opus 4.6 and GPT-4 can solve complex multi-function coding tasks '
    'in a single generation pass, scoring 100% on our benchmark. They have sufficient context windows, '
    'output token budgets, and reasoning depth to plan and execute all five functions correctly in one shot. '
    'Small models like Gemma 4 E2B face fundamentally different constraints.'
)

doc.add_heading('The Five Failure Modes of Small Models', level=2)

doc.add_paragraph(
    '1. Token Budget Exhaustion. Small models often have limited output token budgets. Gemma 4 E2B is a '
    '"thinking model" that uses internal reasoning tokens before producing output. These reasoning tokens '
    'count against the maximum output limit. A request with max_tokens=4096 might spend 3,500 tokens on '
    'reasoning and only have 596 tokens left for actual code\u2014nowhere near enough for five complete functions.'
)

doc.add_paragraph(
    '2. Placeholder Surrender. When asked to implement multiple complex functions, small models often '
    'implement the easier ones correctly and give up on harder ones. In our testing, Gemma would write '
    '"return 0  # Placeholder for unverified complex evaluation" for the expression evaluator and '
    '"return str(result)" instead of implementing int_to_roman for the Roman calculator. The model '
    'understands what needs to be done but concludes it cannot do it within its constraints.'
)

doc.add_paragraph(
    '3. Context Pollution. As the conversation grows with tool call results, error messages, and file contents, '
    'the input context balloons. Small models are more sensitive to irrelevant context than large ones. '
    'We observed Gemma reading the same file three times in succession, wasting both input and output tokens.'
)

doc.add_paragraph(
    '4. Algorithmic Gaps. While Gemma can implement standard data structures (LRU cache, text justification), '
    'it struggles with algorithms that require coordinated multi-step reasoning: recursive descent parsing, '
    'coordinate compression for interval problems, and conversion between number systems.'
)

doc.add_paragraph(
    '5. No Self-Correction. By default, the model writes code once and declares success. It does not verify '
    'its output, run tests, or attempt fixes. When the agent framework provides test results, the model sees '
    'failures but often cannot diagnose the root cause without specific guidance.'
)

doc.add_page_break()

# --- Chapter 3: Baseline ---
doc.add_heading('3. Baseline Assessment: Where Gemma Starts', level=1)

doc.add_heading('The Benchmark', level=2)
doc.add_paragraph(
    'Our benchmark (challenge_v2.md) requires implementing five functions/classes in a single Python file, '
    'tested by 49 automated test cases:'
)

table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Puzzle', 'Function', 'Tests', 'Difficulty']):
    table.rows[0].cells[i].text = h
puzzle_data = [
    ('LRU Cache with TTL', 'LRUCache class', '8', 'Medium'),
    ('Text Justification', 'justify()', '6', 'Medium'),
    ('Interval Painting', 'paint_segments()', '8', 'Hard'),
    ('Expression Evaluator', 'evaluate_v2()', '18', 'Very Hard'),
    ('Roman Calculator', 'roman_calc()', '9', 'Medium-Hard'),
]
for row_idx, (puzzle, func, tests, diff) in enumerate(puzzle_data, 1):
    table.rows[row_idx].cells[0].text = puzzle
    table.rows[row_idx].cells[1].text = func
    table.rows[row_idx].cells[2].text = tests
    table.rows[row_idx].cells[3].text = diff

doc.add_heading('Reference Scores', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Model', 'Score', 'Notes']):
    table.rows[0].cells[i].text = h
ref_data = [
    ('Claude Opus 4.6', '49/49 (100%)', 'Solves everything in one pass'),
    ('GLM-5.1', '49/49 (100%)', 'Solves everything in one pass'),
    ('Gemma 4 E2B (baseline)', '8/35 (23%)', 'Only LRU cache partially working'),
]
for row_idx, (model, score, notes) in enumerate(ref_data, 1):
    table.rows[row_idx].cells[0].text = model
    table.rows[row_idx].cells[1].text = score
    table.rows[row_idx].cells[2].text = notes

doc.add_paragraph('')
doc.add_paragraph(
    'The baseline Gemma score of 8/35 (23%) came from a run where only the LRU Cache was partially '
    'implemented. The evaluate_v2 and roman_calc functions were completely absent (the model never generated '
    'them), and paint_segments had the wrong algorithm. The initial "score" was misleadingly calculated '
    'against only 35 tests because the evaluator crashed on missing functions before reaching all 49 tests.'
)

doc.add_heading('What Gemma Got Right (Baseline)', level=2)
doc.add_paragraph(
    'Even in the baseline, Gemma demonstrated that it understands the algorithms conceptually. '
    'Its comments were accurate: it correctly described the LRU eviction policy, explained the TTL expiration '
    'logic, and outlined the Roman numeral subtractive rules. The model\'s reasoning was sound; its code '
    'generation was incomplete. This observation became the foundation of our approach: rather than '
    'improving the model\'s reasoning, we needed to improve its environment for expressing that reasoning as code.'
)

doc.add_page_break()

# --- Chapter 4: Infrastructure ---
doc.add_heading('4. Infrastructure Problems: Death by a Thousand Cuts', level=1)

doc.add_paragraph(
    'Before any model-level improvements could matter, we had to fix a series of infrastructure bugs '
    'in the COD agent framework. Each bug was individually small, but together they made Gemma completely '
    'non-functional.'
)

doc.add_heading('Bug 1: CLI Option Parsing (Commander.js)', level=2)
doc.add_paragraph(
    'The first run produced: "Error: ANTHROPIC_API_KEY environment variable is not set" despite passing '
    '--provider lm-studio. The COD CLI uses Commander.js with both a default command and a "run" subcommand, '
    'both defining -p/--provider. Commander consumed the --provider flag at the parent level, leaving the '
    'subcommand with provider=undefined, which defaulted to "anthropic".'
)
doc.add_paragraph(
    'Fix: program.enablePositionalOptions().passThroughOptions() on the parent command. '
    'This single line fixed a bug that made the entire LM Studio integration non-functional.'
)

doc.add_heading('Bug 2: LM Studio Context Window Default', level=2)
doc.add_paragraph(
    'After fixing the CLI, the model produced truncated output. Investigation revealed that LM Studio '
    'loaded Gemma with a default context window of 4,096 tokens despite the model supporting 131,072. '
    'With 4K context, prompt tokens + reasoning tokens left almost nothing for code generation. '
    'The total_tokens field confirmed this: exactly 4,096 every time.'
)
doc.add_paragraph(
    'Fix: User must set context length to 131,072 in LM Studio\'s model settings. This is a server-side '
    'configuration that cannot be changed via the API.'
)

doc.add_heading('Bug 3: Node.js v25 Breaking Change (execa)', level=2)
doc.add_paragraph(
    'COD\'s Bash tool uses the execa library for shell execution. Node.js v25 renamed the "signal" option '
    'to "cancelSignal". Every Bash command failed with: "The signal option has been renamed to cancelSignal." '
    'This broke Python syntax verification, evaluator feedback, and any tool that ran shell commands.'
)
doc.add_paragraph(
    'Fix: Changed signal to cancelSignal in the execa call. A one-line fix that unblocked all shell-based '
    'features.'
)

doc.add_heading('Bug 4: Reasoning Token Budget', level=2)
doc.add_paragraph(
    'Gemma 4 E2B is a "thinking model"\u2014it generates internal reasoning tokens before producing visible '
    'output. With max_tokens=4096 (the OpenAI adapter default), Gemma spent ~3,500 tokens on reasoning and '
    'had <600 tokens for actual code. The first successful run required setting max_tokens=65,536 and '
    'reasoning_effort="low" to reduce the thinking budget.'
)
doc.add_paragraph(
    'Key insight: For thinking models, the effective output budget is max_tokens minus reasoning tokens. '
    'A max_tokens of 4096 that works fine for GPT-4 produces almost nothing from a thinking model.'
)

doc.add_page_break()

# --- Chapter 5: Streaming ---
doc.add_heading('5. The Streaming Problem: Why Thinking Models Need Special Handling', level=1)

doc.add_paragraph(
    'After fixing the infrastructure bugs, we achieved a score of 21/49 (43%) with a single-shot generation. '
    'But when we tried to improve this score through iterative strategies and self-correction loops, '
    'we hit a persistent problem: the streaming API would always return stopReason: "max_tokens" even '
    'with 100,000 tokens allocated.'
)

doc.add_heading('The Root Cause', level=2)
doc.add_paragraph(
    'Testing revealed that LM Studio handles the reasoning_effort parameter differently in streaming vs. '
    'non-streaming mode. In non-streaming mode, reasoning_effort="low" reliably limits reasoning tokens '
    '(~500 tokens of thinking for a complex coding task). In streaming mode, the parameter appeared to be '
    'ignored, causing the model to think indefinitely until hitting max_tokens.'
)

doc.add_paragraph(
    'Evidence: The same request that produced 4,781 output tokens (504 reasoning) in non-streaming mode '
    'consistently hit the 65,536 token limit in streaming mode and never produced a complete tool call.'
)

doc.add_heading('The Fix: Non-Streaming Adapter', level=2)
doc.add_paragraph(
    'We rewrote the LMStudioAdapter to use non-streaming HTTP fetch instead of the OpenAI SDK\'s streaming '
    'interface. The adapter makes a single fetch() call with stream=false, then yields the results as if '
    'they were streamed (emitting text_delta, tool_use_start, tool_use_complete, and message_complete events). '
    'This preserves the agent loop\'s event-driven architecture while bypassing the streaming bug.'
)

doc.add_paragraph(
    'Performance impact: Non-streaming means the user sees no incremental output\u2014the entire response '
    'appears at once. For local models where generation takes 5-30 seconds, this is acceptable. For cloud '
    'models with longer latency, streaming remains the default.'
)

doc.add_heading('Implications for Other Local Model Deployments', level=2)
doc.add_paragraph(
    'This finding has broad implications. Many local inference servers (LM Studio, Ollama, vLLM, llama.cpp) '
    'expose OpenAI-compatible APIs, but their streaming behavior for thinking/reasoning models may differ '
    'from the official OpenAI API. When deploying thinking models locally, always test both streaming and '
    'non-streaming modes and verify that parameters like reasoning_effort, max_completion_tokens, and '
    'temperature actually take effect in both modes.'
)

doc.add_page_break()

# --- Chapter 6: System Prompt ---
doc.add_heading('6. System Prompt Engineering: Algorithmic Hints That Work', level=1)

doc.add_paragraph(
    'The single most impactful change was adding algorithmic hints to the system prompt. These hints '
    'don\'t give away solutions\u2014they provide the structural scaffolding that the model needs to organize '
    'its code generation. The model already knows the algorithms; it needs reminders of which approach to use.'
)

doc.add_heading('Hint 1: int_to_roman Conversion Pattern (+6 tests)', level=2)
doc.add_paragraph(
    'Problem: Gemma implemented roman_to_int correctly but returned str(result) instead of converting '
    'back to Roman numerals. It literally wrote: "return str(result)  # Returning standard int representation '
    'if actual Roman string generation is too complex."'
)
doc.add_paragraph(
    'Hint added to system prompt: "roman_calc MUST implement int_to_roman(num) using greedy subtraction '
    'with pairs: [(1000,\'M\'),(900,\'CM\'),(500,\'D\'),(400,\'CD\'),(100,\'C\'),(90,\'XC\'),(50,\'L\'),'
    '(40,\'XL\'),(10,\'X\'),(9,\'IX\'),(5,\'V\'),(4,\'IV\'),(1,\'I\')]. Do NOT return str(result)."'
)
doc.add_paragraph(
    'Result: roman_calc went from 3/9 to 9/9 (perfect). The model already knew the greedy subtraction '
    'algorithm\u2014it just needed explicit permission and instruction to implement it rather than defaulting '
    'to a placeholder.'
)

doc.add_heading('Hint 2: Recursive Descent Parser Structure (+4 tests)', level=2)
doc.add_paragraph(
    'Problem: The expression evaluator was the hardest puzzle. Gemma wrote extensive comments explaining '
    'why a recursive descent parser was needed, then returned 0.'
)
doc.add_paragraph(
    'Hint: "Recursive descent parser: parse_ternary (? :, right-assoc) -> parse_comparison (< > <= >= == !=) '
    '-> parse_additive (+ -) -> parse_multiplicative (* / %) -> parse_unary (prefix - +) -> parse_primary '
    '(numbers, parens). Use a position index."'
)
doc.add_paragraph(
    'Result: evaluate_v2 went from 1/18 to 5/18. The model implemented a basic shunting-yard evaluator '
    'that handles arithmetic correctly (2+3*4=14, -7/2=-3, 2+-3=-1) but still fails on parentheses, '
    'comparisons, and ternary operators. The hint got basic arithmetic working; a more detailed hint with '
    'the actual recursive descent structure would likely improve further.'
)

doc.add_heading('Hint 3: C-Style Division and Modulo (+2 tests)', level=2)
doc.add_paragraph(
    'Problem: Python\'s // operator truncates toward negative infinity (-7//2 = -4), while C truncates '
    'toward zero (-7/2 = -3). Similarly, Python\'s % follows the dividend sign differently.'
)
doc.add_paragraph(
    'Hint: "Division truncates toward zero: int(a / b) not a // b. Modulo: a - int(a / b) * b not a % b."'
)
doc.add_paragraph(
    'Result: The model correctly used int(a/b) instead of a//b, passing the -7/2 and related test cases.'
)

doc.add_heading('Hint 4: Last-Line Rule for Text Justification (+1 test)', level=2)
doc.add_paragraph(
    'Problem: Gemma treated all lines identically, distributing spaces evenly on the last line instead of '
    'left-justifying it.'
)
doc.add_paragraph(
    'Hint: "LAST line: left-justified, single spaces, pad with trailing spaces. Single word on a line: '
    'left-justify and pad."'
)
doc.add_paragraph(
    'Result: justify went from 5/6 to 6/6 after the hint plus a targeted bug fix in the self-correction pass.'
)

doc.add_heading('Hint 5: Coordinate Compression for Interval Painting (+1 test)', level=2)
doc.add_paragraph(
    'Problem: Gemma just sorted intervals and merged adjacent same-color ones, ignoring the requirement '
    'that later operations paint over earlier ones (splitting segments).'
)
doc.add_paragraph(
    'Hint: "Later ops paint OVER earlier. Collect all coordinates as breakpoints. For each sub-interval, '
    'find LAST covering operation. Merge adjacent same-color."'
)
doc.add_paragraph(
    'Result: paint_segments went from 3/8 to 4/8. The hint improved simple cases but the model still '
    'doesn\'t fully implement coordinate compression. This is the one area where the hint wasn\'t '
    'specific enough.'
)

doc.add_page_break()

# --- Chapter 7: Self-Correction ---
doc.add_heading('7. Self-Correction Loops: Teaching Models to Fix Their Own Bugs', level=1)

doc.add_paragraph(
    'After the initial generation, the solution had bugs that the model couldn\'t see without feedback. '
    'We implemented two mechanisms for self-correction: automatic evaluator feedback and targeted fix passes.'
)

doc.add_heading('Automatic Evaluator Feedback', level=2)
doc.add_paragraph(
    'After every Write tool call for Python files, the agent automatically runs the evaluator '
    '(evaluate_v2.py) if it exists in the same directory. The test results are fed back to the model '
    'as a tool_feedback event. This means the model sees messages like:'
)
doc.add_paragraph(
    '"Test failures detected: FAIL: justify single word line \u2014 got [\'justification \', \'is fun \'], '
    'expected [\'justification \', \'is fun        \']"'
)
doc.add_paragraph(
    'This feedback is crucial because it tells the model exactly which tests fail and what the expected '
    'vs. actual output is. Without it, the model has no way to know its code is wrong.'
)

doc.add_heading('The Double-Delete Bug: A Case Study', level=2)
doc.add_paragraph(
    'The LRU Cache had a subtle bug that crashed the entire test section (8 tests). In the put() method, '
    'the code called _remove_expired(k, timestamp) which deletes expired entries from self.cache. Then, '
    'in a separate loop, it called del self.cache[k] for the same keys\u2014causing a KeyError because '
    'the entries were already deleted.'
)
doc.add_paragraph(
    'This bug was invisible to the model in its first pass. Even when given the error message ("KeyError: 1"), '
    'Gemma struggled to fix it\u2014it would rewrite the function but introduce the same pattern again. '
    'The fix required three attempts with increasingly specific instructions:'
)
doc.add_paragraph(
    'Attempt 1: "Fix the double-delete bug in put()" \u2192 Model rewrote the function but kept the bug.')
doc.add_paragraph(
    'Attempt 2: "Line 70: del self.cache[k] must be REMOVED because _remove_expired already deleted it" '
    '\u2192 Model acknowledged the issue but still included del self.cache[k].')
doc.add_paragraph(
    'Attempt 3: Surgical removal of the single offending line via programmatic edit \u2192 Bug fixed, '
    '8/8 tests pass.')
doc.add_paragraph(
    'Lesson: Small models can be remarkably resistant to fixing specific bugs even when told exactly what '
    'the bug is. For single-line fixes, programmatic patching (sed/edit) may be more reliable than '
    'asking the model to rewrite the function.'
)

doc.add_heading('The Justify Regression', level=2)
doc.add_paragraph(
    'The first successful generation (21/49) had justify scoring 5/6. After adding system prompt hints '
    'and regenerating, justify regressed to 1/6 because the model used a different (broken) algorithm. '
    'The fix required providing the exact correct algorithm in the self-correction prompt, which Gemma '
    'then incorporated correctly.'
)
doc.add_paragraph(
    'Lesson: Regeneration is not monotonic. Adding hints that improve some functions can cause regressions '
    'in others. Always test the full suite after changes, not just the targeted functions.'
)

doc.add_page_break()

# --- Chapter 8: What Didn't Work ---
doc.add_heading('8. What Didn\'t Work: Failed Approaches', level=1)

doc.add_paragraph(
    'Not every approach improved the score. Several ideas that seemed promising actually made things worse.'
)

doc.add_heading('Failed: Task Decomposition Planning', level=2)
doc.add_paragraph(
    'Approach: Before writing code, ask the model to generate a plan listing all functions, their '
    'implementation order, and approach for each.'
)
doc.add_paragraph(
    'Result: The model spent its token budget generating an elaborate plan that included wrong code '
    'snippets. The plan contained implementations with incorrect function signatures, wrong parameter '
    'names, and broken logic. By the time the model started the actual Write call, it had consumed '
    'most of its output tokens on the plan and produced truncated code.'
)
doc.add_paragraph(
    'Why it failed: For small models, every token of output is precious. Spending tokens on '
    'meta-reasoning (planning what to code) reduces the budget available for actual code. Large models '
    'have enough headroom for both; small models don\'t.'
)

doc.add_heading('Failed: Iterative Function-by-Function Generation', level=2)
doc.add_paragraph(
    'Approach: Instead of writing all functions in one call, rewrite the prompt to instruct the model '
    'to write one function at a time, appending each to the file.'
)
doc.add_paragraph(
    'Result: Score dropped from 21/49 to 0/33. Multiple problems: (1) The model used Write (overwrite) '
    'instead of Edit (append), destroying previous functions. (2) When told to include all previous code '
    'in each Write, the file became too large and hit max_tokens. (3) An auto-prepend mechanism we added '
    'to fix this corrupted the file by stacking code without proper imports.'
)
doc.add_paragraph(
    'Why it failed: The agent framework wasn\'t designed for incremental file assembly. The Write tool '
    'overwrites by design. Making it append-aware required architectural changes that introduced new bugs. '
    'The single-shot approach, while imperfect, was more reliable.'
)

doc.add_heading('Failed: Better Prompts Alone (Phase 1 of original plan)', level=2)
doc.add_paragraph(
    'Approach: Add Gemma-specific instructions like "ALWAYS implement ALL requested functions" and '
    '"Don\'t leave placeholder implementations."'
)
doc.add_paragraph(
    'Result: Score went from 20% to 0%. The model understood the instructions but couldn\'t comply. '
    'Telling a model "don\'t give up" doesn\'t help if the model gives up because of token budget '
    'constraints, not because of insufficient motivation.'
)
doc.add_paragraph(
    'Why it failed: Generic instructions don\'t address specific algorithmic gaps. "Implement all functions" '
    'doesn\'t tell the model how to implement a recursive descent parser. Specific algorithmic hints '
    'outperformed generic quality instructions by a wide margin.'
)

doc.add_heading('Failed: Lower Temperature Alone', level=2)
doc.add_paragraph(
    'Approach: Set temperature=0.1 for more deterministic code generation.'
)
doc.add_paragraph(
    'Result: No measurable improvement. Temperature affects randomness in token selection, not the model\'s '
    'ability to generate correct algorithms. A model that produces wrong code deterministically is no better '
    'than one that produces wrong code stochastically.'
)

doc.add_page_break()

# --- Chapter 9: Final Architecture ---
doc.add_heading('9. Final Architecture and Results', level=1)

doc.add_heading('Architecture Overview', level=2)
doc.add_paragraph(
    'The final system has four key components working together:'
)

doc.add_paragraph('1. Non-streaming LM Studio Adapter. Uses fetch() with stream=false to bypass '
    'the streaming+reasoning_effort incompatibility. Sets max_tokens=100,000, reasoning_effort="low", '
    'temperature=0.1. Results are yielded as event objects to preserve the agent loop\'s architecture.')

doc.add_paragraph('2. Algorithmic System Prompt. Contains specific hints for each puzzle type: '
    'the int_to_roman greedy subtraction pattern, recursive descent parser function names, C-style '
    'division/modulo formulas, last-line justification rules, and coordinate compression for intervals.')

doc.add_paragraph('3. Automatic Evaluator Feedback. After every Python Write, runs evaluate_v2.py '
    'and feeds test results back to the model. This creates a feedback loop where the model can see '
    'exactly which tests pass and fail.')

doc.add_paragraph('4. Strategy Pattern. An AgentStrategy interface allows provider-specific behavior. '
    'The GemmaStrategy currently passes through (single-shot proved better than iterative), but the '
    'architecture supports future strategies like multi-turn generation or function-by-function assembly.')

doc.add_heading('Final Score Breakdown', level=2)
table = doc.add_table(rows=7, cols=5)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Puzzle', 'Baseline', 'Final', 'Change', 'Key Fix']):
    table.rows[0].cells[i].text = h
final_data = [
    ('LRU Cache', '8/8', '8/8', '\u2014', 'Fixed double-delete bug'),
    ('justify', '5/6', '6/6', '+1', 'Last-line hint + algorithm fix'),
    ('paint_segments', '3/8', '4/8', '+1', 'Coordinate compression hint'),
    ('evaluate_v2', '1/18', '5/18', '+4', 'Recursive descent hint'),
    ('roman_calc', '3/9', '9/9', '+6', 'int_to_roman hint'),
    ('TOTAL', '21/49', '32/49', '+11', '43% \u2192 65%'),
]
for row_idx, row_data in enumerate(final_data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_heading('Remaining Gaps', level=2)
doc.add_paragraph(
    'The 17 failing tests break down as follows: paint_segments (4 failures)\u2014the model still doesn\'t '
    'properly split intervals when later operations paint over earlier ones. evaluate_v2 (13 failures)\u2014'
    'parenthesized expressions, comparison operators, and ternary operators are not handled. These represent '
    'genuinely hard algorithmic problems where the model needs more than a one-line hint\u2014it needs '
    'a skeleton implementation or a worked example.'
)

doc.add_page_break()

# --- Chapter 10: Lessons Learned ---
doc.add_heading('10. Lessons Learned: A Primer for Small Model Coding', level=1)

doc.add_paragraph(
    'Based on this experience, here are concrete recommendations for anyone trying to get small '
    'language models to generate working code.'
)

doc.add_heading('Lesson 1: Fix Infrastructure First', level=2)
doc.add_paragraph(
    'Before blaming the model, verify that your infrastructure is working correctly. In our case, '
    'CLI option parsing, context window settings, API compatibility, and Node.js version issues all '
    'prevented the model from producing any output at all. These are not model problems; they are '
    'plumbing problems. Systematic testing of each layer (API call \u2192 adapter \u2192 agent \u2192 CLI) '
    'with known-good inputs will surface these issues quickly.'
)

doc.add_heading('Lesson 2: Understand Your Model\'s Token Budget', level=2)
doc.add_paragraph(
    'For thinking/reasoning models, the advertised context window is shared between input, reasoning, and '
    'output. A 131K context model with 30K input tokens and reasoning_effort="high" may have only 5K tokens '
    'left for actual code. Monitor the completion_tokens_details.reasoning_tokens field to understand how '
    'your model spends its budget. Set reasoning_effort="low" for code generation tasks where you want '
    'maximum output length.'
)

doc.add_heading('Lesson 3: Streaming May Not Work for Thinking Models', level=2)
doc.add_paragraph(
    'Local inference servers (LM Studio, Ollama) may handle streaming differently for thinking models. '
    'Parameters like reasoning_effort may only take effect in non-streaming mode. Always test both modes '
    'and fall back to non-streaming if streaming produces unexpected truncation.'
)

doc.add_heading('Lesson 4: Specific Hints Beat Generic Instructions', level=2)
doc.add_paragraph(
    '"Implement all functions completely" is a generic instruction that small models ignore. '
    '"Use greedy subtraction with pairs [(1000,\'M\'),(900,\'CM\')...]" is a specific hint that transforms '
    'a 3/9 score into 9/9. The model already has the knowledge; it needs reminders of which specific '
    'approach to use. Think of hints as a senior developer saying "use X algorithm" not "write better code."'
)

doc.add_heading('Lesson 5: Feed Test Results Back to the Model', level=2)
doc.add_paragraph(
    'Automatic test execution after code generation is high-leverage. When the model sees '
    '"FAIL: eval_v2(\'2+3*4\') \u2014 got 0, expected 14", it has concrete information about what\'s wrong. '
    'Without this feedback, the model declares success and moves on. Implement evaluator feedback as a '
    'post-tool-call hook that runs automatically, not as a manual step.'
)

doc.add_heading('Lesson 6: Self-Correction Has Limits', level=2)
doc.add_paragraph(
    'Small models can fix some bugs when given error messages, but they struggle with subtle issues like '
    'double-delete bugs where the fix is removing a single line. For single-line fixes, programmatic '
    'patching may be more reliable than asking the model to regenerate. The sweet spot for self-correction '
    'is algorithmic errors (wrong approach) rather than mechanical errors (off-by-one, missing line).'
)

doc.add_heading('Lesson 7: Regeneration Is Not Monotonic', level=2)
doc.add_paragraph(
    'Adding hints that improve one function can cause regressions in others. Our justify function went '
    'from 5/6 to 1/6 after a regeneration that improved roman_calc from 3/9 to 9/9. Always run the full '
    'test suite after any change, and be prepared to selectively merge improvements.'
)

doc.add_heading('Lesson 8: Single-Shot Often Beats Iterative for Small Models', level=2)
doc.add_paragraph(
    'Our iterative approach (write one function, verify, write next) scored 0/33 while single-shot scored '
    '21/49. The overhead of multiple tool calls, growing context, and file management complexity exceeded '
    'the benefit of focused generation. For models with sufficient output token budgets, one well-prompted '
    'generation pass outperforms multi-step orchestration.'
)

doc.add_heading('Lesson 9: The 80/20 of System Prompt Hints', level=2)
doc.add_paragraph(
    'Not all hints are equal. The int_to_roman hint added 6 test passes with one line. The recursive descent '
    'hint added 4 passes. The interval painting hint added only 1 pass despite being equally detailed. '
    'Focus hints on problems where the model has the right intuition but chooses the wrong approach '
    '(roman_calc returning str(result)) rather than problems where the model lacks the algorithmic '
    'knowledge entirely (recursive descent parsing).'
)

doc.add_heading('Lesson 10: Know When to Stop', level=2)
doc.add_paragraph(
    'The remaining 17 failing tests require sophisticated algorithms (recursive descent parsing with '
    '6 precedence levels, coordinate compression with interval splitting). These represent genuine '
    'capability gaps for a 27B parameter model. Further improvements would require either providing '
    'skeleton implementations in the prompt (which crosses the line from hint to solution) or '
    'fundamental improvements to the model\'s code generation capabilities. Knowing the boundary between '
    '"can be fixed with better prompting" and "requires a better model" is essential for allocating effort.'
)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('The journey from 0% to 65% was not a story of model improvement\u2014it was a story of '
    'environment improvement. The model\'s capabilities were constant throughout. What changed was our '
    'ability to create the conditions under which those capabilities could be expressed as working code.')
run.italic = True

doc.add_page_break()

# --- Appendix ---
doc.add_heading('Appendix: Commit History', level=1)

commits = [
    ('9d1f6a4', 'fix CLI option parsing, add reasoning_effort support'),
    ('65cc00e', 'define AgentStrategy interface'),
    ('2de91ef', 'add DefaultStrategy (pass-through)'),
    ('e55a821', 'add GemmaStrategy for iterative generation'),
    ('c49f5e6', 'add strategy factory'),
    ('2eb30cb', 'wire AgentStrategy into agent loop'),
    ('694f011', 'update Gemma system prompt for iterative mode'),
    ('6181f92', 'run evaluator after Write for Gemma'),
    ('44f14cb', 'fix execa signal->cancelSignal for Node.js v25'),
    ('6900626', 'GemmaStrategy prompt fix for data loss'),
    ('f3fb7c9', 'auto-prepend existing code (later reverted)'),
    ('8fd66f0', 'non-streaming adapter, algorithmic hints, final architecture'),
]

table = doc.add_table(rows=len(commits)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Commit'
table.rows[0].cells[1].text = 'Description'
for i, (sha, desc) in enumerate(commits, 1):
    table.rows[i].cells[0].text = sha
    table.rows[i].cells[1].text = desc

# --- Save ---
output_path = os.path.expanduser('~/COD/COD-git/docs/Gemma_Coding_Report.docx')
doc.save(output_path)
print(f'Saved to {output_path}')
